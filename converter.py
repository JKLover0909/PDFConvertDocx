import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import threading
import os
import io
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import shutil


class PDFConverterApp:
    def __init__(self, root):
        self.root = root
        self.root.title("PDF → DOCX Converter")
        self.root.geometry("480x320")
        self.root.resizable(False, False)
        self.root.configure(bg="#f0f4f8")
        self._build_ui()
        self._setup_drop(root)

    def _build_ui(self):
        tk.Label(
            self.root, text="PDF → DOCX Converter",
            font=("Segoe UI", 16, "bold"), bg="#f0f4f8", fg="#1a202c"
        ).pack(pady=(24, 4))

        tk.Label(
            self.root, text="Kéo thả file PDF vào đây hoặc bấm để chọn file",
            font=("Segoe UI", 10), bg="#f0f4f8", fg="#4a5568"
        ).pack()

        self.drop_frame = tk.Frame(
            self.root, width=380, height=100,
            bg="#e2e8f0", relief="flat", bd=2, cursor="hand2"
        )
        self.drop_frame.pack(pady=16)
        self.drop_frame.pack_propagate(False)

        self.drop_label = tk.Label(
            self.drop_frame, text="📄  Chọn file PDF",
            font=("Segoe UI", 11), bg="#e2e8f0", fg="#718096", cursor="hand2"
        )
        self.drop_label.place(relx=0.5, rely=0.5, anchor="center")

        self.drop_frame.bind("<Button-1>", lambda e: self._browse_file())
        self.drop_label.bind("<Button-1>", lambda e: self._browse_file())

        self.progress = ttk.Progressbar(self.root, mode="indeterminate", length=380)

        self.status_var = tk.StringVar(value="")
        tk.Label(
            self.root, textvariable=self.status_var,
            font=("Segoe UI", 10), bg="#f0f4f8", fg="#4a5568", wraplength=420
        ).pack(pady=(0, 8))

        self.btn = tk.Button(
            self.root, text="Convert",
            font=("Segoe UI", 11, "bold"),
            bg="#3182ce", fg="white", activebackground="#2b6cb0",
            relief="flat", padx=32, pady=8, cursor="hand2",
            command=self._start_convert, state="disabled"
        )
        self.btn.pack()
        self.pdf_path = None

    def _setup_drop(self, root):
        try:
            from tkinterdnd2 import DND_FILES
            self.drop_frame.drop_target_register(DND_FILES)
            self.drop_frame.dnd_bind("<<Drop>>", self._on_drop)
            self.drop_label.drop_target_register(DND_FILES)
            self.drop_label.dnd_bind("<<Drop>>", self._on_drop)
        except Exception:
            pass

    def _on_drop(self, event):
        path = event.data.strip().strip("{}")
        if path.lower().endswith(".pdf"):
            self._set_file(path)
        else:
            messagebox.showwarning("Lỗi", "Vui lòng chọn file PDF!")

    def _browse_file(self):
        path = filedialog.askopenfilename(
            title="Chọn file PDF", filetypes=[("PDF files", "*.pdf")]
        )
        if path:
            self._set_file(path)

    def _set_file(self, path):
        self.pdf_path = path
        self.drop_label.config(text=f"📄  {os.path.basename(path)}", fg="#2d3748")
        self.drop_frame.config(bg="#bee3f8")
        self.drop_label.config(bg="#bee3f8")
        self.btn.config(state="normal")
        self.status_var.set("")

    def _start_convert(self):
        if not self.pdf_path:
            return
        self.btn.config(state="disabled")
        self.progress.pack(pady=(0, 8))
        self.progress.start(10)
        self.status_var.set("Đang convert, vui lòng đợi…")
        threading.Thread(target=self._convert, daemon=True).start()

    # ------------------------------------------------------------------ #
    # Core conversion: text stays text, images stay images
    # ------------------------------------------------------------------ #

    def _convert(self):
        tmp_dir = tempfile.mkdtemp()
        try:
            out_path = os.path.splitext(self.pdf_path)[0] + ".docx"
            doc = Document()

            # Remove default empty paragraph
            for para in doc.paragraphs:
                _remove_paragraph(para)

            pdf = fitz.open(self.pdf_path)
            page_width_inch = pdf[0].rect.width / 72 if len(pdf) > 0 else 8.27

            # Set document page size to match first page
            section = doc.sections[0]
            section.page_width = int(page_width_inch * 914400)
            section.page_height = int((pdf[0].rect.height / 72) * 914400)
            section.left_margin = int(0.5 * 914400)
            section.right_margin = int(0.5 * 914400)
            section.top_margin = int(0.5 * 914400)
            section.bottom_margin = int(0.5 * 914400)

            usable_width = page_width_inch - 1.0  # account for margins

            for page_num, page in enumerate(pdf):
                if page_num > 0:
                    doc.add_page_break()

                # Get all blocks sorted top→bottom, left→right
                blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
                blocks.sort(key=lambda b: (round(b["bbox"][1] / 10), b["bbox"][0]))

                for block in blocks:
                    if block["type"] == 0:
                        # TEXT block
                        self._add_text_block(doc, block)
                    elif block["type"] == 1:
                        # IMAGE block — use the embedded image bytes directly
                        self._add_image_block(doc, block, usable_width, tmp_dir)

            pdf.close()
            doc.save(out_path)
            self.root.after(0, self._on_success, out_path)
        except Exception as e:
            self.root.after(0, self._on_error, str(e))
        finally:
            shutil.rmtree(tmp_dir, ignore_errors=True)

    def _add_text_block(self, doc, block):
        for line in block["lines"]:
            para = doc.add_paragraph()
            for span in line["spans"]:
                text = span["text"]
                if not text.strip():
                    continue
                run = para.add_run(text)
                flags = span.get("flags", 0)
                run.bold = bool(flags & 16)    # bit 4 = bold
                run.italic = bool(flags & 2)   # bit 1 = italic
                size = span.get("size", 11)
                run.font.size = Pt(max(6, min(size, 72)))
                # Preserve text color
                color = span.get("color", 0)
                if color and color != 0:
                    r = (color >> 16) & 0xFF
                    g = (color >> 8) & 0xFF
                    b = color & 0xFF
                    run.font.color.rgb = RGBColor(r, g, b)

    def _add_image_block(self, doc, block, max_width_inch, tmp_dir):
        img_bytes = block.get("image")
        if not img_bytes:
            return
        ext = block.get("ext", "png")
        img_path = os.path.join(tmp_dir, f"img_{id(block)}.{ext}")
        with open(img_path, "wb") as f:
            f.write(img_bytes)

        bbox = block["bbox"]
        img_width_inch = (bbox[2] - bbox[0]) / 72
        display_width = min(img_width_inch, max_width_inch)

        para = doc.add_paragraph()
        run = para.add_run()
        try:
            run.add_picture(img_path, width=Inches(display_width))
        except Exception:
            # Fallback: convert to PNG via PyMuPDF if format unsupported
            pix = fitz.Pixmap(io.BytesIO(img_bytes))
            png_path = img_path + ".png"
            pix.save(png_path)
            run.add_picture(png_path, width=Inches(display_width))

    def _on_success(self, out_path):
        self.progress.stop()
        self.progress.pack_forget()
        self.status_var.set(f"✅ Đã lưu: {os.path.basename(out_path)}")
        self.btn.config(state="normal")
        messagebox.showinfo("Thành công", f"Đã convert xong!\n\nFile lưu tại:\n{out_path}")

    def _on_error(self, err):
        self.progress.stop()
        self.progress.pack_forget()
        self.status_var.set("❌ Lỗi khi convert!")
        self.btn.config(state="normal")
        messagebox.showerror("Lỗi", f"Không thể convert file:\n{err}")


def _remove_paragraph(para):
    """Remove a paragraph element from its parent."""
    p = para._element
    p.getparent().remove(p)


def main():
    try:
        from tkinterdnd2 import TkinterDnD
        root = TkinterDnD.Tk()
    except Exception:
        root = tk.Tk()
    app = PDFConverterApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()

    def __init__(self, root):
        self.root = root
        self.root.title("PDF → DOCX Converter")
        self.root.geometry("480x320")
        self.root.resizable(False, False)
        self.root.configure(bg="#f0f4f8")

        self._build_ui()
        self._setup_drop(root)

    def _build_ui(self):
        tk.Label(
            self.root, text="PDF → DOCX Converter",
            font=("Segoe UI", 16, "bold"), bg="#f0f4f8", fg="#1a202c"
        ).pack(pady=(24, 4))

        tk.Label(
            self.root, text="Kéo thả file PDF vào đây hoặc bấm để chọn file",
            font=("Segoe UI", 10), bg="#f0f4f8", fg="#4a5568"
        ).pack()

        self.drop_frame = tk.Frame(
            self.root, width=380, height=100,
            bg="#e2e8f0", relief="flat", bd=2,
            cursor="hand2"
        )
        self.drop_frame.pack(pady=16)
        self.drop_frame.pack_propagate(False)

        self.drop_label = tk.Label(
            self.drop_frame, text="📄  Chọn file PDF",
            font=("Segoe UI", 11), bg="#e2e8f0", fg="#718096",
            cursor="hand2"
        )
        self.drop_label.place(relx=0.5, rely=0.5, anchor="center")

        self.drop_frame.bind("<Button-1>", lambda e: self._browse_file())
        self.drop_label.bind("<Button-1>", lambda e: self._browse_file())

        self.progress = ttk.Progressbar(
            self.root, mode="indeterminate", length=380
        )

        self.status_var = tk.StringVar(value="")
        self.status_label = tk.Label(
            self.root, textvariable=self.status_var,
            font=("Segoe UI", 10), bg="#f0f4f8", fg="#4a5568",
            wraplength=420
        )
        self.status_label.pack(pady=(0, 8))

        self.btn = tk.Button(
            self.root, text="Convert",
            font=("Segoe UI", 11, "bold"),
            bg="#3182ce", fg="white", activebackground="#2b6cb0",
            relief="flat", padx=32, pady=8, cursor="hand2",
            command=self._start_convert, state="disabled"
        )
        self.btn.pack()

        self.pdf_path = None

    def _setup_drop(self, root):
        try:
            from tkinterdnd2 import DND_FILES
            self.drop_frame.drop_target_register(DND_FILES)
            self.drop_frame.dnd_bind("<<Drop>>", self._on_drop)
            self.drop_label.drop_target_register(DND_FILES)
            self.drop_label.dnd_bind("<<Drop>>", self._on_drop)
        except Exception:
            pass

    def _on_drop(self, event):
        path = event.data.strip().strip("{}")
        if path.lower().endswith(".pdf"):
            self._set_file(path)
        else:
            messagebox.showwarning("Lỗi", "Vui lòng chọn file PDF!")

    def _browse_file(self):
        path = filedialog.askopenfilename(
            title="Chọn file PDF",
            filetypes=[("PDF files", "*.pdf")]
        )
        if path:
            self._set_file(path)

    def _set_file(self, path):
        self.pdf_path = path
        name = os.path.basename(path)
        self.drop_label.config(text=f"📄  {name}", fg="#2d3748")
        self.drop_frame.config(bg="#bee3f8")
        self.drop_label.config(bg="#bee3f8")
        self.btn.config(state="normal")
        self.status_var.set("")

    def _start_convert(self):
        if not self.pdf_path:
            return
        self.btn.config(state="disabled")
        self.progress.pack(pady=(0, 8))
        self.progress.start(10)
        self.status_var.set("Đang convert, vui lòng đợi…")
        threading.Thread(target=self._convert, daemon=True).start()

    def _convert(self):
        tmp_dir = tempfile.mkdtemp()
        try:
            out_path = os.path.splitext(self.pdf_path)[0] + ".docx"
            doc = Document()

            pdf = fitz.open(self.pdf_path)
            for i, page in enumerate(pdf):
                # Render page at 200 DPI for crisp image quality
                mat = fitz.Matrix(200 / 72, 200 / 72)
                pix = page.get_pixmap(matrix=mat, alpha=False)

                img_path = os.path.join(tmp_dir, f"page_{i}.png")
                pix.save(img_path)

                # Set page orientation to match PDF page
                width_inch = page.rect.width / 72
                height_inch = page.rect.height / 72

                section = doc.sections[-1] if i == 0 else doc.add_section()
                section.page_width = int(width_inch * 914400)   # EMU
                section.page_height = int(height_inch * 914400)
                section.left_margin = 0
                section.right_margin = 0
                section.top_margin = 0
                section.bottom_margin = 0

                # Add page break between pages (except first)
                if i > 0:
                    para = doc.paragraphs[-1] if doc.paragraphs else doc.add_paragraph()
                else:
                    para = doc.add_paragraph()

                run = para.add_run()
                run.add_picture(img_path, width=Inches(width_inch))

            pdf.close()
            doc.save(out_path)
            self.root.after(0, self._on_success, out_path)
        except Exception as e:
            self.root.after(0, self._on_error, str(e))
        finally:
            # Clean up temp images
            import shutil
            shutil.rmtree(tmp_dir, ignore_errors=True)

    def _on_success(self, out_path):
        self.progress.stop()
        self.progress.pack_forget()
        self.status_var.set(f"✅ Đã lưu: {os.path.basename(out_path)}")
        self.btn.config(state="normal")
        messagebox.showinfo("Thành công", f"Đã convert xong!\n\nFile lưu tại:\n{out_path}")

    def _on_error(self, err):
        self.progress.stop()
        self.progress.pack_forget()
        self.status_var.set("❌ Lỗi khi convert!")
        self.btn.config(state="normal")
        messagebox.showerror("Lỗi", f"Không thể convert file:\n{err}")


def main():
    try:
        from tkinterdnd2 import TkinterDnD
        root = TkinterDnD.Tk()
    except Exception:
        root = tk.Tk()
    app = PDFConverterApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
